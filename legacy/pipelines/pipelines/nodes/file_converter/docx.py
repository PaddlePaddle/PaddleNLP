# Copyright (c) 2022 PaddlePaddle Authors. All Rights Reserved.
# Copyright 2021 deepset GmbH. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import hashlib
import logging
import os
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

import docx
from docx.document import Document
from docx.oxml.shape import CT_Picture
from docx.parts.image import ImagePart
from docx.text.paragraph import Paragraph
from PIL import Image
from pipelines.nodes.file_converter import BaseConverter

logger = logging.getLogger(__name__)


class DocxToTextConverter(BaseConverter):
    def __init__(
        self,
        remove_numeric_tables: bool = False,
        valid_languages: Optional[List[str]] = None,
    ):
        """
        :param remove_numeric_tables: This option uses heuristics to remove numeric rows from the tables.
                                      The tabular structures in documents might be noise for the reader model if it
                                      does not have table parsing capability for finding answers. However, tables
                                      may also have long strings that could possible candidate for searching answers.
                                      The rows containing strings are thus retained in this option.
        :param valid_languages: validate languages from a list of languages specified in the ISO 639-1
                                (https://en.wikipedia.org/wiki/ISO_639-1) format.
                                This option can be used to add test for encoding errors. If the extracted text is
                                not one of the valid languages, then it might likely be encoding error resulting
                                in garbled text.
        """

        # Save init parameters to enable export of component config as YAML
        self.set_config(remove_numeric_tables=remove_numeric_tables, valid_languages=valid_languages)

        self.remove_numeric_tables = remove_numeric_tables
        self.valid_languages = valid_languages

        self.desc_path = "parse_files"
        os.makedirs(self.desc_path, exist_ok=True)

    def convert(
        self,
        file_path: Path,
        meta: Optional[Dict[str, Any]] = None,
        remove_numeric_tables: Optional[bool] = None,
        valid_languages: Optional[List[str]] = None,
        encoding: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        """
        Extract text from a .docx file.
        Note: As docx doesn't contain "page" information, we actually extract and return a list of paragraphs here.
        For compliance with other converters we nevertheless opted for keeping the methods name.

        :param file_path: Path to the .docx file you want to convert
        :param meta: dictionary of meta data key-value pairs to append in the returned document.
        :param remove_numeric_tables: This option uses heuristics to remove numeric rows from the tables.
                                      The tabular structures in documents might be noise for the reader model if it
                                      does not have table parsing capability for finding answers. However, tables
                                      may also have long strings that could possible candidate for searching answers.
                                      The rows containing strings are thus retained in this option.
        :param valid_languages: validate languages from a list of languages specified in the ISO 639-1
                                (https://en.wikipedia.org/wiki/ISO_639-1) format.
                                This option can be used to add test for encoding errors. If the extracted text is
                                not one of the valid languages, then it might likely be encoding error resulting
                                in garbled text.
        :param encoding: Not applicable
        """
        if remove_numeric_tables is None:
            remove_numeric_tables = self.remove_numeric_tables
        if valid_languages is None:
            valid_languages = self.valid_languages
        if remove_numeric_tables is True:
            raise Exception("'remove_numeric_tables' is not supported by DocxToTextConverter.")
        if valid_languages is True:
            raise Exception("Language validation using 'valid_languages' is not supported by DocxToTextConverter.")
        # Creating word reader object.
        file = docx.Document(file_path)
        documents = []
        text_dict = {}
        # This part will parse the docs files with images, the text and the following images will be added as an document
        for i in range(len(file.paragraphs)):
            paragraph = file.paragraphs[i]
            # Extracting images from the paragraph
            image_list = self.get_image_list(file, paragraph)
            # Extracting text from the paragraph
            # If there is text, Adding the text to text_dict
            if paragraph.text != "":
                text = paragraph.text
                if bool(text_dict) is False:
                    text_dict = {"text": [text], "images": []}
                else:
                    text_dict["text"].append(text)
                if image_list is not None:
                    image_names = self.save_images(image_list)
                    text_dict["images"] += image_names
            else:
                # If there are not text and images, adding text_dict to documents
                if image_list is None and bool(text_dict):
                    raw_text = "".join(text_dict["text"])
                    # If the extracted text is "", skip it
                    if raw_text == "":
                        continue
                    meta_data = {}
                    if meta is not None and "name" in meta:
                        meta_data["name"] = meta["name"]
                    meta_data["images"] = text_dict["images"]
                    document = {"content": raw_text, "content_type": "text", "meta": meta_data}
                    documents.append(document)

                    text = paragraph.text
                    text_dict = {"text": [text], "images": []}
                elif image_list is not None:
                    image_names = self.save_images(image_list)
                    text_dict["images"] += image_names
                else:
                    continue
        return documents

    def save_images(self, image_list):
        """
        Save the parsed image into desc_path
        :param image_list: image files from the docx file
        """
        image_names = []
        for i, image in enumerate(image_list):
            if image:
                # File extension & file content
                ext, blob = image.ext, image.blob
                # Using md5 to generate image name and save image into desc_path
                md5hash = hashlib.md5(blob)
                md5_name = md5hash.hexdigest()
                image_name = "{}_{}.{}".format(md5_name, i, ext)
                image_path = os.path.join(self.desc_path, image_name)
                Image.open(BytesIO(blob)).save(image_path)
                # Adding image_name into the text_dict as the image for the text
                image_names.append(image_name)

        return image_names

    def get_image_list(self, document: Document, paragraph: Paragraph):
        """
        Extract images from  paragraph and document object.
        :param document: file objects
        :param paragraph: image paragraph
        """
        result_list = []
        # Looking up the images of the paragraph
        img_list = paragraph._element.xpath(".//pic:pic")
        if len(img_list) == 0 or not img_list:
            return
        # Extracting images from the document
        for i in range(len(img_list)):
            img: CT_Picture = img_list[i]
            embed = img.xpath(".//a:blip/@r:embed")[0]
            related_part: ImagePart = document.part.related_parts[embed]
            image: Image = related_part.image
            result_list.append(image)
        return result_list


class DocxTotxtConverter(BaseConverter):
    def convert(
        self,
        file_path: Path,
        separator="\n",
        **kwargs: Any,
    ) -> List[str]:
        """
        Extract text from a .docx file.
        """
        # Creating word reader object.
        file = docx.Document(file_path)
        txt_documents = ""
        txt_documents = separator.join([i.text for i in file.paragraphs])
        document = {"content": txt_documents, "content_type": "text", "meta": {}}
        return [document]
